"""Word document export service"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Optional
from datetime import datetime
import io

from app.models.schemas import ActionItem


class WordExportService:
    """Service for generating Word documents"""
    
    def _is_rtl_text(self, text: str) -> bool:
        """
        Check if text contains significant RTL characters (Hebrew, Arabic)
        
        Args:
            text: Text to check
        
        Returns:
            True if text contains >30% RTL characters
        """
        if not text:
            return False
        
        rtl_chars = 0
        # Hebrew Unicode range: U+0590 to U+05FF
        # Arabic Unicode range: U+0600 to U+06FF
        for char in text:
            if ('\u0590' <= char <= '\u05FF') or ('\u0600' <= char <= '\u06FF'):
                rtl_chars += 1
        
        return rtl_chars > len(text) * 0.3  # >30% RTL characters
    
    def _set_rtl_paragraph(self, paragraph):
        """
        Set paragraph direction to RTL (Right-to-Left)
        
        Args:
            paragraph: docx paragraph object
        """
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        # Also set right alignment for RTL text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def create_document(
        self,
        transcription: str,
        summary: str,
        participants: List[str],
        decisions: List[str],
        action_items: List[ActionItem],
        filename: Optional[str] = "meeting_transcription"
    ) -> io.BytesIO:
        """
        Create a Word document with meeting transcription and analysis
        
        Args:
            transcription: Full transcription text
            summary: Meeting summary
            participants: List of participants
            decisions: List of decisions
            action_items: List of action items
            filename: Base filename (without extension)
        
        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()
        
        # Detect if content is RTL (Hebrew/Arabic)
        is_rtl = self._is_rtl_text(transcription + summary)
        
        # Title
        title = doc.add_heading('Meeting Transcription & Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.runs[0].font.size = Pt(10)
        doc.add_paragraph()  # Spacing
        
        # Summary Section
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(summary)
        if is_rtl:
            self._set_rtl_paragraph(summary_para)
        doc.add_paragraph()  # Spacing
        
        # Participants Section
        doc.add_heading('Participants', level=1)
        if participants:
            for participant in participants:
                p = doc.add_paragraph(participant, style='List Bullet')
                if is_rtl:
                    self._set_rtl_paragraph(p)
        else:
            doc.add_paragraph('No participants identified.')
        doc.add_paragraph()  # Spacing
        
        # Decisions Section
        doc.add_heading('Decisions', level=1)
        if decisions:
            for decision in decisions:
                p = doc.add_paragraph(decision, style='List Bullet')
                if is_rtl:
                    self._set_rtl_paragraph(p)
        else:
            doc.add_paragraph('No decisions recorded.')
        doc.add_paragraph()  # Spacing
        
        # Action Items Section
        doc.add_heading('Action Items', level=1)
        if action_items:
            for item in action_items:
                task_para = doc.add_paragraph(f'Task: {item.task}', style='List Bullet')
                if is_rtl:
                    self._set_rtl_paragraph(task_para)
                assignee_para = doc.add_paragraph(f'  Assignee: {item.assignee}', style='List Bullet 2')
                if is_rtl:
                    self._set_rtl_paragraph(assignee_para)
                if item.deadline:
                    deadline_para = doc.add_paragraph(f'  Deadline: {item.deadline}', style='List Bullet 2')
                    if is_rtl:
                        self._set_rtl_paragraph(deadline_para)
        else:
            doc.add_paragraph('No action items identified.')
        doc.add_paragraph()  # Spacing
        
        # Full Transcription Section
        doc.add_heading('Full Transcription', level=1)
        transcription_para = doc.add_paragraph(transcription)
        if is_rtl:
            self._set_rtl_paragraph(transcription_para)
        
        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream

